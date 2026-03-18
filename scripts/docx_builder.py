from pathlib import Path

from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent.parent
DOCX_DIR = BASE_DIR / "docs" / "docx"


LABELS = {
    "fr": ("QR lecteur audio:", "Lien lecteur audio:"),
    "ru": ("QR \u0430\u0443\u0434\u0438\u043e\u043f\u043b\u0435\u0435\u0440\u0430:", "\u0421\u0441\u044b\u043b\u043a\u0430 \u043d\u0430 \u0430\u0443\u0434\u0438\u043e\u043f\u043b\u0435\u0435\u0440:"),
    "en": ("Audio player QR:", "Audio player link:"),
    "ar": ("\u0631\u0645\u0632 QR \u0644\u0645\u0634\u063a\u0644 \u0627\u0644\u0635\u0648\u062a:", "\u0631\u0627\u0628\u0637 \u0645\u0634\u063a\u0644 \u0627\u0644\u0635\u0648\u062a:"),
    "tr": ("Ses oynatici QR kodu:", "Ses oynatici baglantisi:"),
}


SOURCE_IMAGES_TITLES = {
    "fr": "Images du document source",
    "ru": "\u0418\u0437\u043e\u0431\u0440\u0430\u0436\u0435\u043d\u0438\u044f \u0438\u0437 \u0438\u0441\u0445\u043e\u0434\u043d\u043e\u0433\u043e \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430",
    "en": "Images from the source document",
    "ar": "\u0635\u0648\u0631 \u0645\u0646 \u0627\u0644\u0645\u0633\u062a\u0646\u062f \u0627\u0644\u0623\u0635\u0644\u064a",
    "tr": "Kaynak belgedeki gorseller",
}


def labels_pour_langue(langue):
    return LABELS.get((langue or "fr").lower(), LABELS["fr"])


def titre_images_source(langue):
    return SOURCE_IMAGES_TITLES.get((langue or "fr").lower(), SOURCE_IMAGES_TITLES["fr"])


def ajouter_lignes_texte(doc, texte):
    for ligne in texte.splitlines():
        ligne = ligne.strip()
        doc.add_paragraph(ligne if ligne else "")


def ajouter_section(doc, section, index=0):
    titre = section.get("titre", "")
    texte = section.get("texte", "")
    qr_path = section.get("qr_path")
    audio_url = section.get("audio_url")
    langue = section.get("langue", "fr")

    label_qr, label_lien = labels_pour_langue(langue)

    if index > 0:
        doc.add_paragraph("")

    doc.add_heading(titre, level=1)
    ajouter_lignes_texte(doc, texte)

    if qr_path and Path(qr_path).exists():
        doc.add_paragraph(label_qr)
        doc.add_picture(str(qr_path), width=Inches(1.5))

    if audio_url:
        doc.add_paragraph(f"{label_lien} {audio_url}")


def ajouter_images_source(doc, images_source, langue="fr"):
    if not images_source:
        return

    doc.add_page_break()
    doc.add_heading(titre_images_source(langue), level=1)

    for image_path in images_source:
        image_file = Path(image_path)
        if not image_file.exists():
            continue
        doc.add_paragraph(image_file.name)
        doc.add_picture(str(image_file), width=Inches(5.5))


def creer_docx(nom_base, sections, images_source=None):
    """Cree un DOCX editable a partir des sections fournies."""
    try:
        DOCX_DIR.mkdir(parents=True, exist_ok=True)
        docx_path = DOCX_DIR / f"{nom_base}.docx"
        doc = Document()

        for index, section in enumerate(sections):
            ajouter_section(doc, section, index)

        langue_doc = sections[0].get("langue", "fr") if sections else "fr"
        ajouter_images_source(doc, images_source, langue_doc)

        doc.save(str(docx_path))
        print("DOCX cree:", docx_path)
        return str(docx_path)
    except Exception as e:
        print("Erreur creation DOCX:", e)
        return None
